"""Render a practical Markdown subset to DOCX.

Supported: ATX headings, paragraphs, bullet/numbered lists, simple pipe tables,
inline links, and fenced code blocks. Limitations: no nested tables, images,
raw HTML, footnotes, task-list controls, or full CommonMark inline formatting.
Complex Markdown should be rendered with a dedicated Markdown/Pandoc pipeline.
"""

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    raise SystemExit("pip install python-docx  (see requirements.txt)")

HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
BULLET_RE = re.compile(r"^(\s*)[-+*]\s+(.+)$")
NUMBER_RE = re.compile(r"^(\s*)\d+[.)]\s+(.+)$")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)\s]+)(?:\s+\"[^\"]*\")?\)")
TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


def add_hyperlink(paragraph, text, url):
    """Append a clickable external hyperlink to a paragraph."""
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_content(paragraph, text):
    """Add text while preserving Markdown links as DOCX hyperlinks."""
    cursor = 0
    for match in LINK_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        add_hyperlink(paragraph, match.group(1), match.group(2))
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def split_table_row(line):
    stripped = line.strip().strip("|")
    return [cell.strip().replace(r"\|", "|") for cell in re.split(r"(?<!\\)\|", stripped)]


def is_table_start(lines, index):
    return (
        index + 1 < len(lines)
        and "|" in lines[index]
        and bool(TABLE_SEPARATOR_RE.match(lines[index + 1]))
    )


def add_table(document, lines, index):
    rows = [split_table_row(lines[index])]
    index += 2  # Skip Markdown separator row.
    while index < len(lines) and lines[index].strip() and "|" in lines[index]:
        rows.append(split_table_row(lines[index]))
        index += 1

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index in range(column_count):
            value = values[column_index] if column_index < len(values) else ""
            cell_paragraph = table.cell(row_index, column_index).paragraphs[0]
            add_inline_content(cell_paragraph, value)
            if row_index == 0:
                for run in cell_paragraph.runs:
                    run.bold = True
    return index


def ensure_code_style(document):
    styles = document.styles
    try:
        return styles["Code Block"]
    except KeyError:
        style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(9)
        style.paragraph_format.space_after = Pt(0)
        return style


def render(markdown_text):
    document = Document()
    code_style = ensure_code_style(document)
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith(fence):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            paragraph = document.add_paragraph(style=code_style)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            continue

        heading = HEADING_RE.match(line)
        if heading:
            paragraph = document.add_heading(level=len(heading.group(1)))
            add_inline_content(paragraph, heading.group(2))
            index += 1
            continue

        if is_table_start(lines, index):
            index = add_table(document, lines, index)
            continue

        bullet = BULLET_RE.match(line)
        numbered = NUMBER_RE.match(line)
        if bullet or numbered:
            match = bullet or numbered
            level = min(len(match.group(1).expandtabs(2)) // 2, 2)
            base_style = "List Bullet" if bullet else "List Number"
            style = base_style if level == 0 else f"{base_style} {level + 1}"
            paragraph = document.add_paragraph(style=style)
            add_inline_content(paragraph, match.group(2))
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if (
                not candidate.strip()
                or HEADING_RE.match(candidate)
                or BULLET_RE.match(candidate)
                or NUMBER_RE.match(candidate)
                or candidate.strip().startswith(("```", "~~~"))
                or is_table_start(lines, index)
            ):
                break
            paragraph_lines.append(candidate.strip())
            index += 1
        paragraph = document.add_paragraph()
        add_inline_content(paragraph, " ".join(paragraph_lines))

    return document


def parse_args():
    parser = argparse.ArgumentParser(
        description="Convert a supported Markdown subset to DOCX.",
        epilog=(
            "Limitations: no images, raw HTML, footnotes, nested tables, task-list "
            "controls, or complete CommonMark inline formatting."
        ),
    )
    parser.add_argument("input", type=Path, help="UTF-8 Markdown input")
    parser.add_argument("output", type=Path, help="DOCX output path")
    return parser.parse_args()


def main():
    args = parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = render(args.input.read_text(encoding="utf-8"))
    document.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
